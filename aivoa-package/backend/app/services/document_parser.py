"""Turn an uploaded complaint document into plain text.

Deliberately lightweight: the assessment says production-grade OCR isn't
required, so the goal here is to reliably recover text from the formats a QA
team actually receives (emailed PDFs, Word files, pasted email bodies) and to
fail with a clear message on anything else.
"""

from __future__ import annotations

import io
import re
from email import policy
from email.parser import BytesParser

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".eml", ".md", ".csv"}


class UnsupportedDocumentError(ValueError):
    """Raised when we can't get text out of the uploaded file."""


def extract_text(filename: str, data: bytes) -> str:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == ".pdf":
        text = _from_pdf(data)
    elif ext == ".docx":
        text = _from_docx(data)
    elif ext == ".eml":
        text = _from_eml(data)
    elif ext in {".txt", ".md", ".csv"}:
        text = data.decode("utf-8", errors="replace")
    elif ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        raise UnsupportedDocumentError(
            "Image files need OCR, which isn't enabled. "
            "Paste the complaint text instead, or upload a PDF, DOCX, TXT or EML file."
        )
    else:
        raise UnsupportedDocumentError(
            f"Can't read '{filename}'. Supported formats: PDF, DOCX, TXT, EML."
        )

    text = normalise(text)
    if not text.strip():
        raise UnsupportedDocumentError(
            "No readable text found in the document. "
            "If it's a scan, paste the complaint text instead."
        )
    return text


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _from_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Complaint forms are often laid out as tables, so read those too.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_eml(data: bytes) -> str:
    message = BytesParser(policy=policy.default).parsebytes(data)
    header = "\n".join(
        f"{name}: {message[name]}"
        for name in ("From", "To", "Date", "Subject")
        if message[name]
    )
    body_part = message.get_body(preferencelist=("plain", "html"))
    body = body_part.get_content() if body_part else ""
    if body_part is not None and body_part.get_content_type() == "text/html":
        body = re.sub(r"<[^>]+>", " ", body)
    return f"{header}\n\n{body}"


def normalise(text: str) -> str:
    """Collapse the whitespace noise that PDF extraction leaves behind."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
